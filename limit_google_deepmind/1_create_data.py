import os
import json
from docx import Document
from dotenv import load_dotenv
load_dotenv()   # This loads the .env file

# Base folder (where this script is located)
DATA_DIR_SOURCE = os.getenv("DATA_DIR_SOURCE")
SOURCE_CORPUS_FILE=os.getenv("SOURCE_CORPUS_FILE")
# Input corpus file
corpus_file = os.path.join(DATA_DIR_SOURCE,SOURCE_CORPUS_FILE )
output_dir = os.getenv("DATA_DIR_TARGET")

# Ensure output directory exists
os.makedirs(output_dir, exist_ok=True)

# Read corpus and write individual DOCX files
with open(corpus_file, "r", encoding="utf-8") as f:
    for line in f:
        doc = json.loads(line.strip())
        doc_id = doc["_id"].replace(" ", "_")  # replace spaces for safe filenames
        title = doc.get("title", "").strip()
        text = doc.get("text", "").strip()

        file_path = os.path.join(output_dir, f"{doc_id}.txt")

        # Skip if file already exists
        if os.path.exists(file_path):
            print(f"⏩ Skipping {file_path} (already exists)")
            continue

        document = Document()
        if title:
            document.add_heading(title, level=1)
        document.add_paragraph(text)

        document.save(file_path)
        print(f" Saved {file_path}")

print(f"\n🎉 DOCX documents exported to folder: {output_dir}")
